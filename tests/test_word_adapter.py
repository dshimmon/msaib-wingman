# Tests Word document extraction and routing.

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SRC_DIRECTORY = PROJECT_ROOT / "src"

sys.path.insert(0, str(SRC_DIRECTORY))

from document_router import extract_document_units
from word_adapter import extract_word_units


class WordAdapterTests(unittest.TestCase):
    def create_document(self, build_document):
        temporary_directory = tempfile.TemporaryDirectory()
        self.addCleanup(temporary_directory.cleanup)

        file_path = (
            Path(temporary_directory.name)
            / "test-document.docx"
        )
        document = Document()
        build_document(document)
        document.save(file_path)

        return file_path

    def create_sectioned_document(self):
        def build_document(document):
            document.add_paragraph("Introductory text")
            document.add_paragraph("")
            document.add_heading("First Section", level=1)
            document.add_paragraph("First paragraph")

            table = document.add_table(rows=3, cols=3)
            table.rows[0].cells[0].text = "Name"
            table.rows[0].cells[1].text = "Role"
            table.rows[0].cells[2].text = "Status"
            table.rows[1].cells[0].text = "Avery"
            table.rows[1].cells[1].text = "Pilot"
            table.rows[1].cells[2].text = "Ready"

            document.add_heading("Second Section", level=2)
            document.add_paragraph("Second paragraph")
            document.add_paragraph("   ")

        return self.create_document(build_document)

    def test_extracts_section_units_in_document_order(self):
        units = extract_word_units(
            self.create_sectioned_document()
        )

        self.assertEqual(
            units,
            [
                {
                    "heading": None,
                    "location": "Section 1",
                    "text": "Introductory text",
                },
                {
                    "heading": "First Section",
                    "location": "Section 2",
                    "text": (
                        "First Section\n"
                        "First paragraph\n"
                        "Name | Role | Status\n"
                        "Avery | Pilot | Ready"
                    ),
                },
                {
                    "heading": "Second Section",
                    "location": "Section 3",
                    "text": (
                        "Second Section\n"
                        "Second paragraph"
                    ),
                },
            ],
        )

    def test_document_without_headings_produces_one_unit(self):
        def build_document(document):
            document.add_paragraph("")
            document.add_paragraph("Standalone paragraph")

        units = extract_word_units(
            self.create_document(build_document)
        )

        self.assertEqual(
            units,
            [
                {
                    "heading": None,
                    "location": "Section 1",
                    "text": "Standalone paragraph",
                }
            ],
        )

    def test_router_routes_docx_files(self):
        file_path = self.create_sectioned_document()

        self.assertEqual(
            extract_document_units(file_path),
            extract_word_units(file_path),
        )

    def test_router_preserves_existing_unsupported_errors(self):
        with self.assertRaisesRegex(
            ValueError,
            r"Unsupported document type: \.rtf",
        ):
                extract_document_units("notes.rtf")


if __name__ == "__main__":
    unittest.main()
